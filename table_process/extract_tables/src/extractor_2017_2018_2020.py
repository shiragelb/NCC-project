"""
Table extractor for special years 2017, 2018, 2020 using Claude API.
"""

import os
import json
import logging
import pandas as pd
from docx import Document
from anthropic import Anthropic
import time

logger = logging.getLogger(__name__)


class TableExtractor2017_2018_2020:
    """
    Extracts tables from Word documents for years 2017, 2018, 2020.
    These years require Claude API for extraction due to complex format.
    """
    
    def __init__(self, reports_dir="/content/reports", tables_dir="/content/tables"):
        self.reports_dir = reports_dir
        self.tables_dir = tables_dir
        self.encoding = "utf-8-sig"
        
        # Initialize Claude client
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            logger.warning("ANTHROPIC_API_KEY not found in environment variables")
            logger.warning("Set it with: export ANTHROPIC_API_KEY='your-key-here'")
            self.client = None
        else:
            self.client = Anthropic(api_key=api_key)
            logger.info("Claude API client initialized")
        
        # Track statistics
        self.total_cost = 0
        self.total_tokens_in = 0
        self.total_tokens_out = 0
        
    def process_files(self, years=None, chapters=None):
        """
        Process Word documents for special years 2017, 2018, 2020.
        
        Args:
            years: List of years to process
            chapters: List of chapters to process
            
        Returns:
            dict: All extracted table summaries
        """
        if not self.client:
            logger.error("Claude API client not initialized. Please set ANTHROPIC_API_KEY environment variable.")
            return {}
            
        if years is None:
            years = [2017, 2018, 2020]
        if chapters is None:
            chapters = range(1, 16)
            
        all_summaries = {}
        
        for year in years:
            year_dir = os.path.join(self.reports_dir, str(year))
            if not os.path.exists(year_dir):
                logger.warning(f"Year directory not found: {year_dir}")
                continue
                
            for chapter in chapters:
                # Try different filename patterns
                possible_files = [
                    f"{chapter:02d}.docx",
                    f"{chapter}.docx",
                    f"{chapter:02d}_*.docx"  # Pattern matching
                ]
                
                file_path = None
                for pattern in possible_files:
                    if '*' in pattern:
                        # Handle pattern matching
                        import glob
                        matches = glob.glob(os.path.join(year_dir, pattern))
                        if matches:
                            file_path = matches[0]
                            break
                    else:
                        test_path = os.path.join(year_dir, pattern)
                        if os.path.exists(test_path):
                            file_path = test_path
                            break
                
                if not file_path:
                    logger.warning(f"No file found for year {year}, chapter {chapter}")
                    continue
                
                logger.info(f"Processing: {file_path}")
                chapter_summaries = self.extract_tables_with_claude(file_path, year, chapter)
                all_summaries.update(chapter_summaries)
                
                # Small delay to avoid rate limiting
                time.sleep(1)
        
        # Save global summary
        self._save_global_summary(all_summaries)
        
        # Print cost summary
        logger.info(f"Total API cost: ${self.total_cost:.4f}")
        logger.info(f"Total tokens: {self.total_tokens_in:,} in, {self.total_tokens_out:,} out")
        
        return all_summaries
    
    def extract_tables_with_claude(self, docx_path, year, chapter):
        """
        Extract tables from a document using Claude API.
        
        Args:
            docx_path: Path to the Word document
            year: Year of the document
            chapter: Chapter number
            
        Returns:
            dict: Table summaries for this document
        """
        try:
            # Read document text
            doc = Document(docx_path)
            full_text = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    full_text.append(row_text)
            
            document_content = "\n".join(full_text)
            
            # Skip if document is empty
            if len(document_content.strip()) < 100:
                logger.warning(f"Document appears empty or too small: {docx_path}")
                return {}
            
            # Craft the prompt for Claude
            prompt = f"""Analyze this Hebrew document and extract ONLY proper tables (not diagrams).

RULES:
1. ONLY extract tables that have "לוח" followed by a number (like "לוח 1.1:", "לוח 1.2:") 
2. IGNORE anything with "תרשים" - these are diagrams, not tables
3. For each table, extract the table ID, name, column names, and data
4. Look for continuation tables marked with "(המשך)"
5. Clean any RTL markers like {{dir="rtl"}}

Return a valid JSON array with this structure for each table:
{{
  "table_id": "1.1",  
  "table_name": "name after colon",
  "full_header": "full header line",
  "is_continuation": false,
  "base_table_id": "1.1",
  "column_names": ["col1", "col2"],
  "data": [["val1", "val2"]]
}}

DOCUMENT (first 40000 chars):
{document_content[:40000]}

Return ONLY valid JSON array, starting with [ and ending with ]"""
            
            # Make API call
            logger.info(f"  Calling Claude API for year {year}, chapter {chapter}...")
            response = self.client.messages.create(
                model="claude-3-5-sonnet-20241022",  # Using Sonnet 3.5
                max_tokens=8000,
                temperature=0,
                messages=[{"role": "user", "content": prompt}]
            )
            
            # Get response text
            response_text = response.content[0].text
            
            # Try to parse JSON
            try:
                result = json.loads(response_text)
            except json.JSONDecodeError:
                # Try to extract JSON from response
                import re
                match = re.search(r'\[.*\]', response_text, re.DOTALL)
                if match:
                    result = json.loads(match.group(0))
                else:
                    logger.error(f"Failed to parse Claude response for {docx_path}")
                    return {}
            
            # Process and save tables
            summaries = self._process_claude_results(result, year, chapter)
            
            # Update cost tracking
            input_tokens = response.usage.input_tokens
            output_tokens = response.usage.output_tokens
            cost = input_tokens * 0.000003 + output_tokens * 0.000015  # Sonnet 3.5 pricing
            
            self.total_cost += cost
            self.total_tokens_in += input_tokens
            self.total_tokens_out += output_tokens
            
            logger.info(f"  Extracted {len(summaries)} tables, cost: ${cost:.4f}")
            
            return summaries
            
        except Exception as e:
            logger.error(f"Error processing {docx_path}: {e}")
            return {}
    
    def _process_claude_results(self, tables_json, year, chapter):
        """
        Process Claude's JSON response and save tables as CSV files.
        
        Args:
            tables_json: JSON array of tables from Claude
            year: Year of document
            chapter: Chapter number
            
        Returns:
            dict: Table summaries
        """
        # Combine continuation tables first
        combined_tables = {}
        
        for table in tables_json:
            table_id = table.get('table_id', 'unknown')
            base_id = table.get('base_table_id', table_id)
            is_continuation = table.get('is_continuation', False)
            
            if is_continuation:
                if base_id in combined_tables:
                    # Add data to existing table
                    combined_tables[base_id]['data'].extend(table.get('data', []))
                else:
                    # Base not found, treat as new table
                    combined_tables[table_id] = table
            else:
                combined_tables[table_id] = table
        
        # Save tables and create summaries
        summaries = {}
        table_serial = 1
        
        # Create output directory
        output_dir = os.path.join(self.tables_dir, str(year), str(chapter))
        os.makedirs(output_dir, exist_ok=True)
        
        for table_id, table in combined_tables.items():
            # Extract table info
            table_name = table.get('table_name', 'unnamed')
            full_header = table.get('full_header', table_name)
            column_names = table.get('column_names', [])
            data = table.get('data', [])
            
            # Skip empty tables
            if not data:
                continue
            
            # Create identifier
            identifier = f"{table_serial}_{chapter}_{year}"
            
            # Save as CSV
            csv_path = os.path.join(output_dir, f"{identifier}.csv")
            
            # First row is the full header, then columns, then data
            with open(csv_path, 'w', newline='', encoding=self.encoding) as f:
                import csv
                writer = csv.writer(f)
                
                # Write header row (full table name)
                writer.writerow([full_header])
                
                # Write column names
                if column_names:
                    writer.writerow(column_names)
                
                # Write data
                writer.writerows(data)
            
            # Add to summaries
            summaries[identifier] = full_header
            logger.info(f"    Saved table {identifier}: {table_name[:50]}")
            
            table_serial += 1
        
        # Save chapter summary
        if summaries:
            summary_path = os.path.join(output_dir, "summaries.json")
            with open(summary_path, 'w', encoding='utf-8') as f:
                json.dump(summaries, f, ensure_ascii=False, indent=2)
        
        return summaries
    
    def _save_global_summary(self, summaries):
        """Save global summary file."""
        if not summaries:
            return
            
        summary_path = os.path.join(self.tables_dir, "..", "tables_summary.json")
        
        # Load existing summary if it exists
        existing_summaries = {}
        if os.path.exists(summary_path):
            with open(summary_path, 'r', encoding='utf-8') as f:
                existing_summaries = json.load(f)
        
        # Merge with new summaries
        existing_summaries.update(summaries)
        
        # Save updated summary
        with open(summary_path, 'w', encoding='utf-8') as f:
            json.dump(existing_summaries, f, ensure_ascii=False, indent=2)